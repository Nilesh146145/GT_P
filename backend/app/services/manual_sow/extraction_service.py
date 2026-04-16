"""
Document text extraction stub (spec §4–5). Phase 1: local PDF/DOCX text + heuristic chunks.
"""

from __future__ import annotations

import io
import re
import uuid
from dataclasses import dataclass
from typing import Any, Dict, List, Tuple

from app.schemas.manual_sow.enums import ContextDetectionStatus, ExtractionCategory
from app.services.manual_sow.platform_detection import infer_platform_type_from_text


@dataclass
class ExtractionResult:
    items: List[Dict[str, Any]]
    report: Dict[str, Any]
    page_count: int


def _split_paragraphs(text: str) -> List[str]:
    parts = [p.strip() for p in re.split(r"\n\s*\n+", text) if len(p.strip()) > 20]
    if not parts:
        parts = [text.strip()[:2000]] if text.strip() else []
    return parts[:40]


def _guess_category(snippet: str) -> Tuple[ExtractionCategory, int]:
    low = snippet.lower()
    if any(k in low for k in ("budget", "cost", "price", "payment")):
        return ExtractionCategory.budget, 78
    if any(k in low for k in ("timeline", "milestone", "deadline", "schedule")):
        return ExtractionCategory.timeline, 76
    if any(k in low for k in ("compliance", "gdpr", "security policy", "legal")):
        return ExtractionCategory.compliance, 80
    if any(k in low for k in ("shall", "must", "requirement", "functional")):
        return ExtractionCategory.features, 82
    if any(k in low for k in ("risk", "assumption")):
        return ExtractionCategory.risk, 72
    if any(k in low for k in ("user", "stakeholder", "role")):
        return ExtractionCategory.user_context, 74
    if any(k in low for k in ("objective", "goal", "vision")):
        return ExtractionCategory.business_objectives, 85
    return ExtractionCategory.assumptions, 65


def extract_bytes(filename: str, content: bytes) -> ExtractionResult:
    """Parse file bytes; returns synthetic extraction items + intelligence report."""
    name = filename.lower()
    text = ""
    pages = 1

    if name.endswith(".pdf"):
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(content))
            pages = max(1, len(reader.pages))
            text = "\n\n".join((p.extract_text() or "") for p in reader.pages)
        except Exception:
            text = ""
    elif name.endswith(".docx"):
        try:
            import docx

            d = docx.Document(io.BytesIO(content))
            text = "\n\n".join(p.text for p in d.paragraphs if p.text.strip())
            pages = max(1, len(d.paragraphs) // 40 + 1)
        except Exception:
            text = ""
    elif name.endswith(".doc"):
        text = ""
    else:
        text = ""

    if not text.strip():
        text = "No extractable text. Please review and enter scope in commercial details."

    paras = _split_paragraphs(text)
    items: List[Dict[str, Any]] = []
    for i, p in enumerate(paras):
        cat, conf = _guess_category(p)
        items.append(
            {
                "public_id": str(uuid.uuid4()),
                "category": cat.value,
                "text": p[:5000],
                "source_page_number": min(pages, 1 + (i % pages)),
                "source_highlight": p[:200],
                "review_state": "pending",
                "edited_text": None,
                "confidence": conf,
                "is_duplicate": False,
                "duplicate_count": 0,
            }
        )

    # Ensure at least one features item for demo flows
    if not any(i["category"] == ExtractionCategory.features.value for i in items):
        items.insert(
            0,
            {
                "public_id": str(uuid.uuid4()),
                "category": ExtractionCategory.features.value,
                "text": "Deliver the capabilities described in the uploaded document.",
                "source_page_number": 1,
                "source_highlight": "capabilities described",
                "review_state": "pending",
                "edited_text": None,
                "confidence": 70,
                "is_duplicate": False,
                "duplicate_count": 0,
            },
        )

    def _ctx(snippet: str) -> str:
        if len(snippet) < 30:
            return ContextDetectionStatus.ABSENT.value
        return ContextDetectionStatus.PRESENT.value if len(snippet) > 120 else ContextDetectionStatus.PARTIAL.value

    obj_snip = " ".join(i["text"] for i in items if i["category"] == "business_objectives")[:500]
    pain_snip = " ".join(i["text"] for i in items if "pain" in i["text"].lower())[:500]
    user_snip = " ".join(i["text"] for i in items if i["category"] == "user_context")[:500]

    report = {
        "contextDetection": {
            "businessObjectives": _ctx(obj_snip),
            "painPoints": _ctx(pain_snip),
            "userContext": _ctx(user_snip),
        },
        "platformType": infer_platform_type_from_text(text),
        "sectionsFound": len(items),
        "aiConfidence": min(95, 60 + len(items) * 2),
        "gapScore": max(40, 85 - len(items)),
        "ambiguities": max(0, len(items) // 5),
        "sensitiveDataDetected": "possible" if any(k in text.lower() for k in ("ssn", "salary", "password")) else "none",
        "sensitiveDataTypes": [],
        "estimatedReviewTime": f"~{max(10, len(items) * 2)} minutes",
    }

    return ExtractionResult(items=items, report=report, page_count=pages)
