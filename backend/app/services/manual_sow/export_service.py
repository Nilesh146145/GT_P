"""PDF / DOCX / JSON export for manual SOW (spec §13)."""

from __future__ import annotations

import io
import json
from datetime import datetime, timezone
from typing import Any, Dict


def build_json_bundle(sow: Dict[str, Any], sections: list, commercial: Dict[str, Any], stages: list) -> Dict[str, Any]:
    sow_out = {k: v for k, v in sow.items() if not k.startswith("_")}
    return {
        "sow": sow_out,
        "sections": sections,
        "commercial_details": commercial,
        "approval_stages": stages,
        "exported_at": datetime.now(timezone.utc).isoformat().replace("+00:00", "Z"),
    }


def render_pdf_bytes(title: str, sections: list) -> bytes:
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
    except ImportError:
        return b"%PDF-1.4 minimal placeholder"

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    width, height = letter
    y = height - 50
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, y, title[:80])
    y -= 30
    c.setFont("Helvetica", 10)
    for sec in sections[:50]:
        title_s = str(sec.get("title", "Section"))[:120]
        body = str(sec.get("content", ""))[:2000].replace("\n", " ")
        for line in _wrap(body, 90):
            if y < 50:
                c.showPage()
                y = height - 50
                c.setFont("Helvetica", 10)
            c.drawString(50, y, line[:100])
            y -= 12
        y -= 8
    c.save()
    return buf.getvalue()


def _wrap(text: str, w: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    cur = ""
    for word in words:
        if len(cur) + len(word) + 1 > w:
            if cur:
                lines.append(cur)
            cur = word
        else:
            cur = f"{cur} {word}".strip()
    if cur:
        lines.append(cur)
    return lines or [""]


def render_docx_bytes(title: str, sections: list) -> bytes:
    try:
        import docx
    except ImportError:
        return b""

    d = docx.Document()
    d.add_heading(title, 0)
    for sec in sections:
        d.add_heading(str(sec.get("title", "Section")), level=1)
        d.add_paragraph(str(sec.get("content", "")))
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()
